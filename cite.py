import re

def bib_modify(filename):
        '''
        Takes filename of a .bib file as a string, scans the file for all titles and applies correct capitalization to the titles.
        '''
    #capute indeces of titles in a bibfile
    indeces = []
    with open(filename, 'rb') as f:
        content = f.readlines()
        decoded = [i.decode() for i in content]
        for i in range(len(decoded)):
            if decoded[i][0:5] == 'title':
                indeces.append(i)
    #capture title and apply rules
    pattern = re.compile(r'({{)(.*)(}})') 
    newl =[]
    for i in indeces:
        cap_ind = [] 
        old = pattern.findall(decoded[i])[0][1].lower().capitalize()
        l = old.split()
        for j in range(len(l)):
            if l[j][-1] in ':?!.':
                cap_ind.append(j+1)
        for e in cap_ind:
            if len(l)>e:
                l[e] = l[e].capitalize()
        new = 'title = {{' + ' '.join(l) + '}}, \n'
        newl.append(new)

    for i in range(len(indeces)):
        decoded[indeces[i]] =  newl[i]
        
    encoded = [i.encode() for i in decoded]  
    with open(filename, 'wb') as f:
        for i in encoded:
            f.write(i)


def hang_ind(filename):
    '''
    Takes a docx filename out of cwd as string, searches for a paragraph that equals 'References' in that document and adds hanging ident and double spacing to all following paragraphs.
    '''
    #import modules
    from docx import Document
    from docx.shared import Pt, Inches, Length
    #load document 
    doc = Document(filename)
    #number of paragraphs
    n = len(doc.paragraphs)
    #find the reference list 
    for p in range(n):
        if doc.paragraphs[p].text == 'References':
            start = p+1 
    #add hanging ident to all following lines (dont use if reflist is not last part)
    for p in range(start, n):
        doc.paragraphs[p].paragraph_format.left_indent = Pt(36)
        doc.paragraphs[p].paragraph_format.first_line_indent = Inches(-0.5)
        doc.paragraphs[p].paragraph_format.line_spacing = 2.0
    doc.save('hang_ind_{}'.format(filename))
    
    
def rb_hang_ind(filename):
    '''
    Takes a docx filename out of cwd as string, searches for a paragraph that equals 'References' in that document and adds hanging ident and double spacing to all following paragraphs.
    '''
    #import modules
    from docx import Document
    from docx.shared import Pt, Inches, Length
    #load document 
    doc = Document(filename)
    #number of paragraphs
    n = len(doc.paragraphs)
    #find the reference list 
    for p in range(n):
        if doc.paragraphs[p].text == '2d. Literature eferences':
            start = p+1 
        if doc.paragraphs[p].text == '2e. Time Plan':
            end = p
    #add hanging ident to all following lines (dont use if reflist is not last part)
    for p in range(start, end):
        doc.paragraphs[p].paragraph_format.left_indent = Pt(36)
        doc.paragraphs[p].paragraph_format.first_line_indent = Inches(-0.5)
        doc.paragraphs[p].paragraph_format.line_spacing = 2.0

    doc.save('hang_ind_{}'.format(filename))
    


