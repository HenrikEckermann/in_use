#Working with (R)markdown and pandoc etc. for academic writing does not yet (in 2017) allow adding hanging indentation in the reference list. This function adds hanging indentation and double line spacing to the reference list. Without customization this only works if the ref-list ist the last part of your document.

def hang_ind(filename):
    '''
    Takes a docx filename out of cwd as string, searches for a paragraph that equals 'References' in that document and adds hanging ident and double spacing to all following paragraphs.
    '''
    #import modules
    from docx import Document
    from docx.shared import Pt, Inches, Length
    #load document 
    doc = Document(filename)
    #number of paragraphs
    n = len(doc.paragraphs)
    #find the reference list 
    for p in range(n):
        if doc.paragraphs[p].text == 'References':
            start = p+1 
    #add hanging ident to all following lines (dont use if reflist is not last part)
    for p in range(start, n):
        doc.paragraphs[p].paragraph_format.left_indent = Pt(36)
        doc.paragraphs[p].paragraph_format.first_line_indent = Inches(-0.5)
        doc.paragraphs[p].paragraph_format.line_spacing = 2.0

    doc.save('hang_ind_{}'.format(filename))

#customization for a doc that I often have to edit
def rb_hang_ind(filename):
    '''
    Takes a docx filename out of cwd as string, searches for a paragraph that equals 'References' in that document and adds hanging ident and double spacing to all following paragraphs.
    '''
    #import modules
    from docx import Document
    from docx.shared import Pt, Inches, Length
    #load document 
    doc = Document(filename)
    #number of paragraphs
    n = len(doc.paragraphs)
    #find the reference list 
    for p in range(n):
        if doc.paragraphs[p].text == '2d. Literature eferences':
            start = p+1 
        if doc.paragraphs[p].text == '2e. Time Plan':
            end = p
    #add hanging ident to all following lines (dont use if reflist is not last part)
    for p in range(start, end):
        doc.paragraphs[p].paragraph_format.left_indent = Pt(36)
        doc.paragraphs[p].paragraph_format.first_line_indent = Inches(-0.5)
        doc.paragraphs[p].paragraph_format.line_spacing = 2.0

    doc.save('hang_ind_{}'.format(filename))
    


